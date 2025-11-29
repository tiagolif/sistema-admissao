import streamlit as st
from fpdf import FPDF
import smtplib
import tempfile
import os
import io
from docx import Document
from docx.shared import Pt
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email import encoders
from datetime import date

# --- CONSTANTES ---
TABELA_SALARIOS = {
    "Agente de Saneamento Nível I": "R$ 1.850,00", "Agente de Saneamento Nível II": "R$ 2.523,40",
    "Agente de Saneamento Nível III": "A Definir", "Motorista D": "A Definir", "Operador de Máquinas": "A Definir",
    "Secretária": "R$ 3.200,00"
}
SECOES_FORMULARIO = {
    "Dados Pessoais": ["Nome Completo", "Data Nasc", "Estado Civil", "Nome Pai", "Nome Mãe", "Nome Cônjuge", "Filhos <14 anos", "Fone Principal", "Fone Recado"],
    "Endereço": ["CEP", "Rua", "Número", "Complemento", "Bairro", "Cidade", "UF"],
    "Documentação Civil": ["CPF", "RG", "RG Órgão", "RG UF", "RG Emissão", "Título de Eleitor", "Título Zona", "Título Seção", "Reservista", "Reservista Série", "Reservista Categoria"],
    "Documentação Trabalhista e Escolaridade": ["CTPS", "CTPS Série", "CTPS UF", "CTPS Emissão", "PIS", "CNH", "CNH Categoria", "CNH Validade", "Escolaridade", "Situação Escolaridade", "Série/Fase"],
    "Dados da Vaga": ["Cargo Pretendido", "Salário", "Data Admissão", "Horário", "VT", "Contrato Experiência"]
}

# --- CONFIGURAÇÕES E ESTADO DA SESSÃO ---
if 'Escolaridade' not in st.session_state: st.session_state['Escolaridade'] = 'Médio'
if 'Situação Escolaridade' not in st.session_state: st.session_state['Situação Escolaridade'] = 'Completo'
if 'Cargo Pretendido' not in st.session_state: st.session_state['Cargo Pretendido'] = ''
if 'passo_atual' not in st.session_state: st.session_state['passo_atual'] = 1
if 'dados_formulario' not in st.session_state: st.session_state['dados_formulario'] = {}
if 'dados_uploads' not in st.session_state: st.session_state['dados_uploads'] = {}

# --- FUNÇÕES DE LÓGICA E NAVEGAÇÃO ---
def salvar_dados_passo(campos):
    for campo in campos: st.session_state['dados_formulario'][campo] = st.session_state.get(campo)
def callback_proximo(campos):
    salvar_dados_passo(campos); st.session_state['passo_atual'] += 1
def callback_anterior():
    st.session_state['passo_atual'] -= 1

# --- FUNÇÕES DE GERAÇÃO DE DOCUMENTOS ---
def get_valor_campo(dados, secao, campo):
    if secao == "Dados da Vaga":
        if campo == "Cargo Pretendido": return dados.get(campo, "N/A")
        if campo == "Salário": return TABELA_SALARIOS.get(dados.get("Cargo Pretendido"), "_________________")
        return "_________________"
    valor = dados.get(campo, "N/A")
    if isinstance(valor, date): return valor.strftime('%d/%m/%Y')
    return str(valor) if valor else "N/A"

def gerar_pdf(dados_formulario, dados_uploads):
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 14); self.cell(0, 10, 'D.S SERVICOS DE ENGENHARIA E SANEAMENTO LTDA', 0, 1, 'C')
            self.set_font('Arial', 'B', 12); self.cell(0, 10, 'Ficha de Admissão Digital', 0, 1, 'C'); self.ln(5)
    pdf = PDF(); pdf.set_auto_page_break(auto=True, margin=15); pdf.add_page(); pdf.set_margins(10, 10, 10); pdf.set_font("Arial", size=11)
    for nome_secao, campos in SECOES_FORMULARIO.items():
        pdf.set_font("Arial", 'B', 12); pdf.cell(0, 10, f"--- {nome_secao} ---", ln=True, align='L'); pdf.ln(2)
        for campo in campos:
            valor_final = get_valor_campo(dados_formulario, nome_secao, campo)
            texto_limpo = valor_final.encode('latin-1', 'replace').decode('latin-1'); campo_limpo = campo.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_font("Arial", 'B', 11); pdf.cell(60, 8, f"{campo_limpo}:"); pdf.set_font("Arial", '', 11); pdf.multi_cell(0, 8, texto_limpo); pdf.ln(1)
        pdf.ln(5)
    if imagens := {k: v for k, v in dados_uploads.items() if v}:
        pdf.add_page(); pdf.set_font("Arial", 'B', 14); pdf.cell(0, 10, 'Documentos Anexados', 0, 1, 'C'); pdf.ln(10)
        for nome_doc, files in imagens.items():
            file_list = files if isinstance(files, list) else [files]
            for idx, uploaded_file in enumerate(file_list):
                if uploaded_file:
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp: tmp.write(uploaded_file.getvalue())
                        pdf.set_font("Arial", 'B', 12); doc_title = nome_doc.replace("_", " ").title()
                        if len(file_list) > 1: doc_title += f" ({idx + 1}/{len(file_list)})"
                        pdf.cell(0, 10, doc_title.encode('latin-1', 'replace').decode('latin-1'), 0, 1); pdf.image(tmp.name, w=170); pdf.ln(10); os.unlink(tmp.name)
                    except Exception as e: st.error(f"Erro ao processar imagem para PDF {nome_doc}: {e}")
    return bytes(pdf.output(dest='S'))

def gerar_word(dados_formulario, dados_uploads):
    doc = Document()
    run = doc.add_heading('D.S SERVICOS DE ENGENHARIA E SANEAMENTO LTDA', level=1).runs[0]
    run.font.name = 'Calibri'; run.font.size = Pt(14); run.bold = True
    doc.add_heading('Ficha de Admissão Digital', level=2)

    for nome_secao, campos in SECOES_FORMULARIO.items():
        doc.add_heading(nome_secao, level=3)
        table = doc.add_table(rows=len(campos), cols=2)
        table.style = 'Table Grid'
        for i, campo in enumerate(campos):
            table.cell(i, 0).text = f"{campo}:"
            table.cell(i, 1).text = get_valor_campo(dados_formulario, nome_secao, campo)

    if imagens := {k: v for k, v in dados_uploads.items() if v}:
        doc.add_page_break()
        doc.add_heading('Documentos Anexados', level=2)
        for nome_doc, files in imagens.items():
            file_list = files if isinstance(files, list) else [files]
            for idx, uploaded_file in enumerate(file_list):
                if uploaded_file:
                    try:
                        doc_title = nome_doc.replace("_", " ").title()
                        if len(file_list) > 1: doc_title += f" ({idx + 1}/{len(file_list)})"
                        doc.add_heading(doc_title, level=3)
                        uploaded_file.seek(0)
                        doc.add_picture(uploaded_file, width=doc.sections[0].page_width * 0.8)
                    except Exception as e: st.error(f"Erro ao processar imagem para DOCX {nome_doc}: {e}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- FUNÇÃO PARA ENVIAR O E-MAIL ---
def enviar_email(pdf_content, word_content, nome_candidato, dados_imagens):
    senha = None
    try:
        remetente = st.secrets["email"]["remetente"]
        senha = st.secrets["email"]["senha"]
        destinatario = st.secrets["email"]["destinatario"]
        email_copia = st.secrets["email"].get("email_copia") # Leitura do email_copia
        
        msg = MIMEMultipart()
        msg["From"] = remetente
        msg["To"] = destinatario
        if email_copia:
            msg["Bcc"] = email_copia # Adiciona o email_copia como BCC
        msg["Subject"] = f"Nova Admissão: {nome_candidato}"
        msg.attach(MIMEText(f"Olá,\n\nSegue em anexo a ficha de admissão preenchida por {nome_candidato} (PDF e DOCX).\n\nAtenciosamente,\nSistema de Admissão Automático", "plain"))
        
        part_pdf = MIMEBase("application", "octet-stream"); part_pdf.set_payload(pdf_content); encoders.encode_base64(part_pdf)
        part_pdf.add_header("Content-Disposition", f"attachment; filename=Admissao_{nome_candidato.replace(' ', '_')}.pdf")
        msg.attach(part_pdf)

        part_word = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document"); part_word.set_payload(word_content); encoders.encode_base64(part_word)
        part_word.add_header("Content-Disposition", f"attachment; filename=Admissao_{nome_candidato.replace(' ', '_')}.docx")
        msg.attach(part_word)

        if dados_imagens:
            for files in dados_imagens.values():
                for uploaded_file in (files if isinstance(files, list) else [files]):
                    if uploaded_file:
                        uploaded_file.seek(0); img_part = MIMEImage(uploaded_file.read(), name=uploaded_file.name)
                        img_part.add_header('Content-Disposition', f'attachment; filename="{uploaded_file.name}"'); msg.attach(img_part)

        server = smtplib.SMTP("smtp.gmail.com", 587); server.starttls(); server.login(remetente, senha)
        
        # Lista de destinatários para o server.sendmail
        destinatarios_finais = [destinatario]
        if email_copia:
            destinatarios_finais.append(email_copia)
        
        server.sendmail(remetente, destinatarios_finais, msg.as_string()) # Envia para múltiplos destinatários
        server.quit()
        return True, None
    except Exception as e:
        erro_msg = str(e)
        if senha: # Só tenta substituir se a senha foi lida com sucesso
             erro_msg = erro_msg.replace(senha, "****")
        return False, erro_msg

# --- LAYOUT E ROTEAMENTO DO WIZARD ---
st.image('logo.png', width=200)
st.markdown("<h4 style='color: #004a99;'>D.S SERVICOS DE ENGENHARIA E SANEAMENTO LTDA</h4>", unsafe_allow_html=True)
st.markdown("<h2 style='color: #004a99;'>Sistema de Admissão Digital</h2>", unsafe_allow_html=True)
passo = st.session_state['passo_atual']

# ... (código dos passos 1-4 mantido idêntico) ...
if passo == 1:
    st.subheader("Passo 1 de 5: Dados Pessoais e Endereço"); campos = ["Nome Completo", "Data Nasc", "Estado Civil", "Nome Pai", "Nome Mãe", "Nome Cônjuge", "Filhos <14 anos", "Fone Principal", "Fone Recado", "CEP", "Rua", "Número", "Complemento", "Bairro", "Cidade", "UF"]
    for c in campos:
        if c not in st.session_state: st.session_state[c] = st.session_state['dados_formulario'].get(c)
    st.write("**Dados Pessoais**"); c1,c2=st.columns(2); c1.text_input("Nome Completo",key="Nome Completo"); c2.date_input("Data Nasc.",min_value=date(1940,1,1),max_value=date.today(),key="Data Nasc"); c1.text_input("Estado Civil",key="Estado Civil"); c2.number_input("Filhos <14 anos (Qtd)",min_value=0,step=1,key="Filhos <14 anos"); c1.text_input("Nome do Pai",key="Nome Pai"); c2.text_input("Nome da Mãe",key="Nome Mãe"); c1.text_input("Nome Cônjuge (se houver)",key="Nome Cônjuge"); c2.text_input("Fone Principal",key="Fone Principal"); c1.text_input("Fone para Recado",key="Fone Recado")
    st.write("---"); st.write("**Endereço**"); c1,c2,c3=st.columns([1,2,0.5]);c1.text_input("CEP",key="CEP");c2.text_input("Rua / Logradouro",key="Rua");c3.text_input("Número",key="Número"); c1,c2=st.columns(2);c1.text_input("Complemento",key="Complemento");c2.text_input("Bairro",key="Bairro"); c1,c2=st.columns([2,0.5]);c1.text_input("Cidade",key="Cidade");c2.text_input("UF",key="UF")
    st.button("Próximo >", on_click=callback_proximo, args=(campos,))
elif passo == 2:
    st.subheader("Passo 2 de 5: Documentação Civil"); campos = ["CPF", "RG", "RG Órgão", "RG UF", "RG Emissão", "Título de Eleitor", "Título Zona", "Título Seção", "Reservista", "Reservista Série", "Reservista Categoria"]
    for c in campos:
        if c not in st.session_state: st.session_state[c] = st.session_state['dados_formulario'].get(c)
    st.write("**Documentos de Identificação**"); c1,c2=st.columns(2);c1.text_input("CPF",key="CPF");c2.text_input("RG",key="RG"); c1,c2,c3=st.columns(3);c1.text_input("Órgão Emissor do RG",key="RG Órgão");c2.text_input("UF do RG",key="RG UF");c3.date_input("Data Emissão do RG",key="RG Emissão")
    st.write("---"); st.write("**Outros Documentos**"); c1,c2,c3=st.columns(3);c1.text_input("Título de Eleitor (Número)",key="Título de Eleitor");c2.text_input("Zona",key="Título Zona");c3.text_input("Seção",key="Título Seção"); c1,c2,c3=st.columns(3);c1.text_input("Reservista (Número)",key="Reservista");c2.text_input("Série",key="Reservista Série");c3.text_input("Categoria",key="Reservista Categoria")
    c1, c2 = st.columns([1, 1]); c1.button("< Anterior", on_click=callback_anterior); c2.button("Próximo >", on_click=callback_proximo, args=(campos,))
elif passo == 3:
    st.subheader("Passo 3 de 5: Documentação Trabalhista e Escolaridade"); campos = ["CTPS", "CTPS Série", "CTPS UF", "CTPS Emissão", "PIS", "CNH", "CNH Categoria", "CNH Validade", "Escolaridade", "Situação Escolaridade", "Série/Fase"]
    for c in campos:
        if c not in st.session_state: st.session_state[c] = st.session_state['dados_formulario'].get(c)
    st.write("**Documentação Trabalhista**"); c1,c2,c3,c4=st.columns(4);c1.text_input("CTPS (Número)",key="CTPS");c2.text_input("Série",key="CTPS Série");c3.text_input("UF",key="CTPS UF");c4.date_input("Data Emissão",key="CTPS Emissão");st.text_input("PIS",key="PIS")
    st.write("---");st.write("**CNH**"); c1,c2,c3=st.columns(3);c1.text_input("CNH (Número)",key="CNH");c2.text_input("Categoria",key="CNH Categoria");c3.date_input("Validade",key="CNH Validade")
    st.write("---");st.write("**Escolaridade**"); c1,c2=st.columns(2);c1.radio("Grau de Escolaridade",["Nenhum","Básico","Médio","Superior","Pós-Graduação"],key="Escolaridade");c2.radio("Situação",["Completo","Incompleto"],key="Situação Escolaridade")
    if st.session_state.get("Situação Escolaridade") == "Incompleto": st.text_input("Qual Série/Fase/Período?", key="Série/Fase")
    c1, c2 = st.columns([1, 1]); c1.button("< Anterior", on_click=callback_anterior); c2.button("Próximo >", on_click=callback_proximo, args=(campos,))
elif passo == 4:
    st.subheader("Passo 4 de 5: Dados da Vaga"); campos = ["Cargo Pretendido"]
    for c in campos:
        if c not in st.session_state: st.session_state[c] = st.session_state['dados_formulario'].get(c)
    with st.expander("Dados da Vaga", expanded=True):
        st.selectbox("Cargo Pretendido", options=[''] + list(TABELA_SALARIOS.keys()), key="Cargo Pretendido", help="Selecione o cargo para o qual você está se candidatando.")
        st.info("As demais informações da vaga, como salário e horário, serão informadas pelo RH.")
    c1, c2 = st.columns([1, 1]); c1.button("< Anterior", on_click=callback_anterior); c2.button("Próximo >", on_click=callback_proximo, args=(campos,))

elif passo == 5:
    st.subheader("Passo 5 de 5: Upload de Documentos")
    st.markdown("<h4 style='color: #004a99;'>Documentos Obrigatórios</h4>", unsafe_allow_html=True)
    def update_uploads():
        for key in st.session_state:
            if key.startswith("upload_"): st.session_state['dados_uploads'][key.replace("upload_", "")] = st.session_state[key]
    
    c1, c2 = st.columns(2)
    c1.file_uploader("Foto 3x4 (Recente)", type=["jpg", "png"], key="upload_foto_3x4", on_change=update_uploads)
    c2.file_uploader("Carteira de Trabalho (CTPS)", type=["jpg", "png", "pdf"], key="upload_ctps", on_change=update_uploads)
    c1.file_uploader("RG (Identidade - Frente e Verso)", type=["jpg", "png", "pdf"], key="upload_rg", on_change=update_uploads)
    c2.file_uploader("CPF", type=["jpg", "png", "pdf"], key="upload_cpf", on_change=update_uploads)
    c1.file_uploader("Título de Eleitor", type=["jpg", "png", "pdf"], key="upload_titulo_eleitor", on_change=update_uploads)
    c2.file_uploader("Certificado de Reservista (para homens)", type=["jpg", "png", "pdf"], key="upload_reservista", on_change=update_uploads)
    c1.file_uploader("Comprovante de Residência", type=["jpg", "png", "pdf"], key="upload_comp_residencia", on_change=update_uploads)
    c2.file_uploader("Carteira de Habilitação (CNH) (se houver)", type=["jpg", "png", "pdf"], key="upload_cnh", on_change=update_uploads)
    c1.file_uploader("Comprovante de Escolaridade (Diplomas/Certificados)", type=["jpg", "png", "pdf"], key="upload_comp_escolaridade", on_change=update_uploads)
    c2.file_uploader("Certidões de Nasc./Casamento (Pode selecionar vários arquivos)", type=["jpg", "png", "pdf"], key="upload_cert_nasc_casamento", on_change=update_uploads, accept_multiple_files=True)
    c1.file_uploader("Cartão de Vacinação dos Filhos (Pode selecionar vários arquivos)", type=["jpg", "png", "pdf"], key="upload_vacinacao_filhos", on_change=update_uploads, accept_multiple_files=True)
    c2.file_uploader("Comprovante de Frequência Escolar (Pode selecionar vários arquivos)", type=["jpg", "png", "pdf"], key="upload_freq_escolar", on_change=update_uploads, accept_multiple_files=True)
    c1.file_uploader("Exame Médico Admissional", type=["jpg", "png", "pdf"], key="upload_exame_medico", on_change=update_uploads)
    
    col1, col2 = st.columns([1, 1]); col1.button("< Anterior", on_click=callback_anterior)
    if col2.button("✅ FINALIZAR E ENVIAR", type="primary"):
        with st.spinner("Finalizando e enviando seus dados..."):
            nome_bruto = st.session_state['dados_formulario'].get("Nome Completo")
            nome_candidato = nome_bruto.strip() if nome_bruto and isinstance(nome_bruto, str) and nome_bruto.strip() else "Candidato_Sem_Nome"
            
            pdf_bytes = gerar_pdf(st.session_state['dados_formulario'], st.session_state['dados_uploads'])
            word_bytes = gerar_word(st.session_state['dados_formulario'], st.session_state['dados_uploads'])
            
            sucesso, erro = enviar_email(pdf_bytes, word_bytes, nome_candidato, st.session_state['dados_uploads'])
            
            if sucesso:
                st.success("🎉 Admissão enviada com sucesso para o RH!"); st.balloons()
            else:
                st.error(f"Houve um problema ao enviar o e-mail: {erro}"); st.warning("Mas não se preocupe! Baixe sua ficha e envie-a manualmente.")

            st.download_button(label="📥 Baixar Ficha em PDF", data=pdf_bytes, file_name=f"Ficha_Admissao_{nome_candidato.replace(' ', '_')}.pdf", mime="application/pdf")
            st.download_button(label="📥 Baixar Ficha em Word (.docx)", data=word_bytes, file_name=f"Ficha_Admissao_{nome_candidato.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            for key in list(st.session_state.keys()):
                if key != 'passo_atual': del st.session_state[key]
            st.info("O formulário foi reiniciado."); st.session_state.passo_atual = 1
